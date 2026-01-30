"""Document parser for CV extraction from PDF and DOCX files."""

import io
from pathlib import Path
from typing import Optional, Union

from ..utils.logger import AgentLogger


class DocumentParser:
    """
    Parse PDF and DOCX documents to extract text content.
    
    Uses PyMuPDF for PDF and python-docx for DOCX files.
    """
    
    def __init__(self):
        self.logger = AgentLogger("DocumentParser")
    
    def parse(self, file_path: Union[str, Path]) -> str:
        """
        Parse a document and extract text content.
        
        Args:
            file_path: Path to PDF or DOCX file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension == ".pdf":
            return self._parse_pdf(path)
        elif extension in [".docx", ".doc"]:
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def parse_bytes(self, content: bytes, filename: str) -> str:
        """
        Parse document content from bytes.
        
        Args:
            content: Raw file bytes
            filename: Original filename (for extension detection)
            
        Returns:
            Extracted text content
        """
        extension = Path(filename).suffix.lower()
        
        if extension == ".pdf":
            return self._parse_pdf_bytes(content)
        elif extension in [".docx", ".doc"]:
            return self._parse_docx_bytes(content)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _parse_pdf(self, path: Path) -> str:
        """Parse PDF file using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
            
            self.logger.info("Parsing PDF", path=str(path))
            
            text_parts = []
            with fitz.open(path) as doc:
                for page_num, page in enumerate(doc):
                    text = page.get_text()
                    text_parts.append(text)
                    self.logger.debug(f"Parsed page {page_num + 1}")
            
            full_text = "\n\n".join(text_parts)
            self.logger.info(f"Extracted {len(full_text)} characters from PDF")
            
            return full_text
            
        except ImportError:
            raise ImportError("PyMuPDF (fitz) not installed. Run: pip install PyMuPDF")
        except Exception as e:
            self.logger.error(f"PDF parsing failed: {e}")
            raise
    
    def _parse_pdf_bytes(self, content: bytes) -> str:
        """Parse PDF from bytes."""
        try:
            import fitz
            
            self.logger.info("Parsing PDF from bytes")
            
            text_parts = []
            with fitz.open(stream=content, filetype="pdf") as doc:
                for page in doc:
                    text_parts.append(page.get_text())
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            raise ImportError("PyMuPDF (fitz) not installed. Run: pip install PyMuPDF")
    
    def _parse_docx(self, path: Path) -> str:
        """Parse DOCX file using python-docx."""
        try:
            from docx import Document
            
            self.logger.info("Parsing DOCX", path=str(path))
            
            doc = Document(path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            full_text = "\n".join(paragraphs)
            if table_text:
                full_text += "\n\n" + "\n".join(table_text)
            
            self.logger.info(f"Extracted {len(full_text)} characters from DOCX")
            
            return full_text
            
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            self.logger.error(f"DOCX parsing failed: {e}")
            raise
    
    def _parse_docx_bytes(self, content: bytes) -> str:
        """Parse DOCX from bytes."""
        try:
            from docx import Document
            
            self.logger.info("Parsing DOCX from bytes")
            
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            return "\n".join(paragraphs)
            
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
